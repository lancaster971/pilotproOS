"""
RAG Management API Endpoints
Following TODO-MILHENA-EXPERT.md specifications exactly

Enterprise-grade RAG system with:
- Document CRUD operations
- Bulk import/export
- Real-time search with relevance scoring
- Version control and metadata management
- WebSocket integration for live updates
"""

import asyncio
import os
import uuid
import json
import zipfile
import tempfile
import aiofiles
from datetime import datetime
from typing import List, Dict, Any, Optional
from io import BytesIO

from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

# LangChain imports (following official documentation)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings

# Local imports
from app.rag.maintainable_rag import MaintainableRAGSystem, DocumentMetadata, DocumentStatus, get_rag_system
from app.security.masking_engine import MultiLevelMaskingEngine, UserLevel
from app.cache.optimized_embeddings_cache import OptimizedEmbeddingsCache
from app.observability.observability import track_api_request, track_business_value

import logging
logger = logging.getLogger(__name__)

# Initialize router
router = APIRouter(prefix="/api/rag", tags=["RAG Management"])

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def send_personal_message(self, message: str, websocket: WebSocket):
        await websocket.send_text(message)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            try:
                await connection.send_text(message)
            except:
                # Remove disconnected clients
                self.active_connections.remove(connection)

manager = ConnectionManager()

# NOTE: Using singleton from maintainable_rag.py (imported above) to avoid duplicate instances
# This ensures GraphSupervisor and API share the same RAG system instance

# Pydantic models for API requests/responses
class DocumentUploadResponse(BaseModel):
    success: bool
    document_id: str
    message: str
    processing_status: str

class DocumentSearchRequest(BaseModel):
    query: str
    filters: Optional[Dict[str, Any]] = None
    k: int = 5
    user_level: UserLevel = UserLevel.BUSINESS

class DocumentSearchResponse(BaseModel):
    success: bool
    results: List[Dict[str, Any]]
    total_count: int
    query_time_ms: float

class DocumentListResponse(BaseModel):
    success: bool
    documents: List[Dict[str, Any]]
    total_count: int
    page: int
    page_size: int

class RAGStatistics(BaseModel):
    total_documents: int
    total_size_mb: float
    categories: List[str]
    avg_document_size: float
    embedding_status: Dict[str, int]
    last_update: datetime

class BulkImportStatus(BaseModel):
    success: bool
    processed_files: int
    failed_files: int
    total_files: int
    processing_time_seconds: float
    errors: List[str]

# Document processing utilities
def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text from various file formats following official documentation"""
    try:
        if filename.endswith('.pdf'):
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        elif filename.endswith('.docx'):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text

        elif filename.endswith('.txt') or filename.endswith('.md'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif filename.endswith('.html'):
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text()

        else:
            raise ValueError(f"Unsupported file format: {filename}")

    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file {filename}: {str(e)}")

# API Endpoints following TODO-MILHENA-EXPERT.md specifications

@router.post("/documents", response_model=DocumentUploadResponse)
async def upload_documents(
    files: List[UploadFile] = File(...),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    auto_category: bool = Form(True),
    extract_metadata: bool = Form(True),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """
    Upload multiple documents with auto-processing
    Supports: PDF, DOCX, MD, TXT, HTML
    """
    try:
        rag = get_rag_system()
        results = []

        # Parse tags if provided
        tag_list = json.loads(tags) if tags else []

        for file in files:
            # Validate file type
            allowed_extensions = ['.pdf', '.docx', '.txt', '.md', '.html']
            if not any(file.filename.endswith(ext) for ext in allowed_extensions):
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file.filename}"
                )

            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                # Extract text content
                text_content = extract_text_from_file(temp_file_path, file.filename)

                # Create document metadata
                doc_metadata = DocumentMetadata(
                    title=file.filename,
                    source="upload",
                    category=category or ("auto" if auto_category else "general"),
                    tags=tag_list,
                    author="user"
                )

                # Create LangChain document
                document = Document(
                    page_content=text_content,
                    metadata=doc_metadata.dict()
                )

                # Add to RAG system
                result = await rag.create_document(
                    content=text_content,
                    metadata=doc_metadata,
                    auto_chunk=True
                )
                doc_id = result.get("doc_id", str(uuid.uuid4()))

                results.append({
                    "filename": file.filename,
                    "document_id": doc_id,
                    "status": "processed",
                    "size": len(content)
                })

                # Broadcast update to WebSocket clients
                await manager.broadcast(json.dumps({
                    "type": "document_uploaded",
                    "document_id": doc_id,
                    "filename": file.filename
                }))

            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)

        # Track business value (metric_type, value) - NOT async
        track_business_value("documents_uploaded", len(files))

        return DocumentUploadResponse(
            success=True,
            document_id=results[0]["document_id"] if results else "",
            message=f"Successfully processed {len(results)} documents",
            processing_status="completed"
        )

    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Error uploading documents: {e}\n{error_trace}")
        raise HTTPException(status_code=500, detail=f"{str(e)} - Check logs for full traceback")

@router.get("/documents", response_model=DocumentListResponse)
async def list_documents(
    category: Optional[str] = None,
    search: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "created_at",
    sort_order: str = "desc"
):
    """List documents with filtering and pagination"""
    try:
        rag = get_rag_system()

        # Build filters
        filters = {}
        if category:
            filters["category"] = category

        # Get ALL documents from RAG system (limit=10000 to get all chunks, then dedup)
        # NOTE: Must use high limit because ChromaDB returns chunks, not unique docs
        documents = await rag.list_documents(
            category_filter=filters.get("category") if filters else None,
            limit=10000  # Get all chunks, dedup happens in list_documents()
        )

        # Apply simple pagination AFTER dedup
        start = (page - 1) * page_size
        end = start + page_size
        paginated_docs = documents[start:end]

        total_count = len(documents)

        return DocumentListResponse(
            success=True,
            documents=paginated_docs,
            total_count=total_count,
            page=page,
            page_size=page_size
        )

    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/documents/{doc_id}")
async def update_document(
    doc_id: str,
    content: str = Form(...),
    metadata: Optional[str] = Form(None)
):
    """Update document content and metadata"""
    try:
        rag = get_rag_system()

        # Parse metadata if provided
        metadata_dict = json.loads(metadata) if metadata else {}

        # Update document (matches maintainable_rag.py signature)
        result = await rag.update_document(
            doc_id=doc_id,
            new_content=content,
            update_metadata=metadata_dict
        )

        if result.get("status") == "error":
            raise HTTPException(status_code=404, detail=result.get("message", "Document not found"))

        # Broadcast update
        await manager.broadcast(json.dumps({
            "type": "document_updated",
            "document_id": doc_id
        }))

        return {"success": True, "message": "Document updated successfully", "result": result}

    except Exception as e:
        logger.error(f"Error updating document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    soft_delete: bool = True
):
    """Delete document with audit trail"""
    try:
        rag = get_rag_system()

        # Delete document (matches maintainable_rag.py signature)
        result = await rag.delete_document(
            doc_id=doc_id,
            hard_delete=not soft_delete
        )

        if result.get("status") == "error":
            raise HTTPException(status_code=404, detail=result.get("message", "Document not found"))

        # Broadcast deletion
        await manager.broadcast(json.dumps({
            "type": "document_deleted",
            "document_id": doc_id
        }))

        return {"success": True, "message": result.get("message", "Document deleted successfully")}

    except Exception as e:
        logger.error(f"Error deleting document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/search", response_model=DocumentSearchResponse)
async def semantic_search(request: DocumentSearchRequest):
    """Semantic search with relevance scoring"""
    try:
        start_time = datetime.now()
        rag = get_rag_system()

        # Perform semantic search (matches maintainable_rag.py method)
        results = await rag.search(
            query=request.query,
            user_level=request.user_level,
            top_k=request.k,
            filter_metadata=request.filters,
            include_archived=False
        )

        # Calculate query time
        query_time = (datetime.now() - start_time).total_seconds() * 1000

        return DocumentSearchResponse(
            success=True,
            results=results,
            total_count=len(results),
            query_time_ms=query_time
        )

    except Exception as e:
        logger.error(f"Error in semantic search: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/stats", response_model=RAGStatistics)

async def get_rag_statistics():
    """RAG system statistics and analytics"""
    try:
        rag = get_rag_system()
        stats = await rag.get_statistics()

        # Extract categories (dict → list)
        categories_dict = stats.get("categories", {})
        categories_list = list(categories_dict.keys()) if isinstance(categories_dict, dict) else []

        # Calculate size estimation (avg 1KB per chunk)
        total_chunks = stats.get("total_chunks", 0)
        total_size_mb = (total_chunks * 1.0) / 1024  # Rough estimate

        # Build embedding status from cache stats
        cache_stats = stats.get("cache_stats", {})
        embedding_status = {
            "cached": cache_stats.get("cache_hits", 0),
            "generated": cache_stats.get("cache_misses", 0),
            "total": cache_stats.get("total_embeddings", 0)
        }

        return RAGStatistics(
            total_documents=stats.get("unique_documents", 0),
            total_size_mb=round(total_size_mb, 2),
            categories=categories_list,
            avg_document_size=round(total_chunks / max(stats.get("unique_documents", 1), 1), 2),
            embedding_status=embedding_status,
            last_update=stats.get("last_update") if stats.get("last_update") else datetime.now()
        )

    except Exception as e:
        logger.error(f"Error getting RAG statistics: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/bulk-import", response_model=BulkImportStatus)

async def bulk_import(archive: UploadFile = File(...)):
    """Bulk import from ZIP archive"""
    try:
        start_time = datetime.now()

        if not archive.filename.endswith('.zip'):
            raise HTTPException(status_code=400, detail="Only ZIP archives are supported")

        rag = get_rag_system()
        processed_files = 0
        failed_files = 0
        errors = []

        # Save uploaded ZIP temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as temp_zip:
            content = await archive.read()
            temp_zip.write(content)
            temp_zip_path = temp_zip.name

        try:
            # Extract and process ZIP contents
            with zipfile.ZipFile(temp_zip_path, 'r') as zip_file:
                total_files = len([f for f in zip_file.namelist() if not f.endswith('/')])

                for file_info in zip_file.infolist():
                    if file_info.is_dir():
                        continue

                    try:
                        # Extract file
                        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                            temp_file.write(zip_file.read(file_info.filename))
                            temp_file_path = temp_file.name

                        try:
                            # Process file
                            text_content = extract_text_from_file(temp_file_path, file_info.filename)

                            # Create document metadata
                            doc_metadata = DocumentMetadata(
                                title=file_info.filename,
                                source="bulk_import",
                                category="imported",
                                author="bulk_import"
                            )

                            # Add to RAG (matches maintainable_rag.py method)
                            result = await rag.create_document(
                                content=text_content,
                                metadata=doc_metadata,
                                auto_chunk=True
                            )

                            if result.get("status") == "success":
                                processed_files += 1

                            # Broadcast progress
                            await manager.broadcast(json.dumps({
                                "type": "bulk_import_progress",
                                "processed": processed_files,
                                "total": total_files,
                                "current_file": file_info.filename
                            }))

                        finally:
                            os.unlink(temp_file_path)

                    except Exception as e:
                        failed_files += 1
                        errors.append(f"{file_info.filename}: {str(e)}")
                        logger.error(f"Error processing {file_info.filename}: {e}")

        finally:
            os.unlink(temp_zip_path)

        processing_time = (datetime.now() - start_time).total_seconds()

        # Track business value (metric_type, value) - NOT async
        track_business_value("bulk_documents_imported", processed_files)

        return BulkImportStatus(
            success=True,
            processed_files=processed_files,
            failed_files=failed_files,
            total_files=processed_files + failed_files,
            processing_time_seconds=processing_time,
            errors=errors[:10]  # Limit error list
        )

    except Exception as e:
        logger.error(f"Error in bulk import: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/reindex")
async def reindex_knowledge_base():
    """
    Force re-indexing of all documents
    NOTE: Currently reindexing is automatic when documents are updated.
    This endpoint triggers a full statistics recalculation.
    """
    try:
        rag = get_rag_system()

        # Get current statistics (triggers internal recalculation)
        stats = await rag.get_statistics()

        # Broadcast reindex completion
        await manager.broadcast(json.dumps({
            "type": "reindex_completed",
            "timestamp": datetime.now().isoformat(),
            "total_documents": stats.get("unique_documents", 0),
            "total_chunks": stats.get("total_chunks", 0)
        }))

        return {
            "success": True,
            "message": "Knowledge base statistics refreshed",
            "status": "completed",
            "stats": stats
        }

    except Exception as e:
        logger.error(f"Error refreshing statistics: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# WebSocket endpoint for real-time updates
@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket for real-time RAG updates"""
    await manager.connect(websocket)
    try:
        while True:
            # Keep connection alive
            data = await websocket.receive_text()
            # Echo received message (can be used for heartbeat)
            await websocket.send_text(f"Received: {data}")
    except WebSocketDisconnect:
        manager.disconnect(websocket)

# Health check endpoint
@router.get("/health")
async def rag_health_check():
    """RAG system health check"""
    try:
        rag = get_rag_system()
        stats = await rag.get_statistics()

        return {
            "status": "healthy",
            "rag_system": "operational",
            "total_documents": stats.get("total_documents", 0),
            "last_check": datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"RAG health check failed: {e}")
        raise HTTPException(status_code=503, detail="RAG system unhealthy")